from __future__ import annotations

from pathlib import Path
from uuid import uuid4

from docx import Document as DocxDocument

from core.exceptions import ParserError
from domain.document import (
    Document,
    DocumentMetadata,
    DocumentType,
    Page,
)
from infrastructure.parsers.base_parser import BaseParser


class DOCXParser(BaseParser):
    """
    Parser for Microsoft Word (.docx) documents.
    """

    @property
    def supported_extensions(self) -> tuple[str, ...]:
        return (".docx",)

    def parse(self, file_path: Path) -> Document:
        """
        Parse a DOCX document into a Document model.
        """

        if not file_path.exists():
            raise ParserError(f"File not found: {file_path}")

        try:
            doc = DocxDocument(file_path)

            paragraphs = [
                paragraph.text
                for paragraph in doc.paragraphs
                if paragraph.text.strip()
            ]

            text = "\n".join(paragraphs)

            metadata = DocumentMetadata(
                title=doc.core_properties.title or "",
                author=doc.core_properties.author or "",
                subject=doc.core_properties.subject or "",
                keywords=[],
                language="",
            )

            page = Page(
                number=1,
                text=text,
            )

            return Document(
                id=str(uuid4()),
                filename=file_path.name,
                filepath=file_path,
                document_type=DocumentType.DOCX,
                pages=[page],
                metadata=metadata,
                extracted_text=text,
                page_count=1,
                file_size=file_path.stat().st_size,
            )

        except Exception as error:
            raise ParserError(
                f"Failed to parse DOCX '{file_path.name}'."
            ) from error